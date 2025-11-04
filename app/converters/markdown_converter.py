"""
Markdown Converter
Location: /mnt/c/Users/Joseph/Documents/Code/md-converter/app/converters/markdown_converter.py

This module provides the core conversion engine for transforming markdown files
with YAML front matter to Word (DOCX) and PDF formats.

Dependencies:
    - python-frontmatter: YAML front matter parsing
    - pypandoc: Markdown to Word conversion
    - weasyprint: PDF generation
    - markdown: HTML rendering

Used by: app/api/routes.py for handling conversion requests
"""
import frontmatter
import pypandoc
import markdown
import logging
from pathlib import Path
from typing import Dict, Tuple, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import WeasyPrint conditionally (only needed for PDF)
try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False
    logging.warning(f"WeasyPrint not available: {e}")


logger = logging.getLogger(__name__)


class ConversionError(Exception):
    """Raised when document conversion fails"""
    pass


class FrontMatterError(Exception):
    """Raised when front matter parsing fails"""
    pass


class MarkdownConverter:
    """
    Converts markdown files with YAML front matter to Word and PDF formats.

    This class handles:
    - Parsing YAML front matter
    - Converting markdown to Word (DOCX) via pypandoc
    - Converting markdown to PDF via weasyprint
    - Applying proper styling and page numbers
    """

    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize converter with optional Word template.

        Args:
            template_path: Path to .docx template with page numbers

        Raises:
            RuntimeError: If Pandoc is not available
        """
        self.template_path = template_path
        self._verify_pandoc()
        logger.info('MarkdownConverter initialized')

    def _verify_pandoc(self) -> None:
        """
        Verify Pandoc is available in the system.

        Raises:
            RuntimeError: If Pandoc is not found
        """
        try:
            version = pypandoc.get_pandoc_version()
            logger.debug(f'Pandoc version {version} detected')
        except OSError as e:
            error_msg = (
                'Pandoc not found. '
                'Install with: apt-get install pandoc or pip install pypandoc-binary'
            )
            logger.error(error_msg)
            raise RuntimeError(error_msg) from e

    def parse_markdown(self, content: str) -> Tuple[Dict, str]:
        """
        Parse markdown with YAML front matter.

        Args:
            content: Raw markdown string with optional YAML front matter

        Returns:
            Tuple of (metadata dict, content string)

        Raises:
            FrontMatterError: If YAML parsing fails critically

        Example:
            >>> metadata, body = converter.parse_markdown(md_string)
            >>> print(metadata['title'])
            'My Document'
        """
        try:
            post = frontmatter.loads(content)
            logger.debug(f'Parsed front matter: {post.metadata.keys()}')
            return post.metadata, post.content
        except Exception as e:
            logger.warning(f'Failed to parse front matter: {e}. Using empty metadata.')
            # If no front matter or parsing fails, return empty dict and full content
            return {}, content

    def format_front_matter(self, metadata: Dict) -> str:
        """
        Format front matter for display in document.

        Args:
            metadata: Dictionary of front matter fields

        Returns:
            Formatted markdown string for document header

        Example:
            >>> formatted = converter.format_front_matter({
            ...     'title': 'Report',
            ...     'author': 'Jane'
            ... })
            >>> print(formatted)
            # Report
            **Author:** Jane
            ---
        """
        if not metadata:
            return ""

        lines = []

        # Title as H1
        title = metadata.get('title', 'Untitled Document')
        lines.append(f"# {title}\n")

        # Common metadata fields
        if 'author' in metadata:
            lines.append(f"**Author:** {metadata['author']}  ")
        if 'date' in metadata:
            date_value = metadata['date']
            # Handle date objects
            if hasattr(date_value, 'strftime'):
                date_value = date_value.strftime('%Y-%m-%d')
            lines.append(f"**Date:** {date_value}  ")
        if 'tags' in metadata:
            tags = metadata['tags']
            if isinstance(tags, list):
                tags = ', '.join(str(t) for t in tags)
            lines.append(f"**Tags:** {tags}  ")

        # Other custom fields
        excluded = {'title', 'author', 'date', 'tags'}
        for key, value in metadata.items():
            if key not in excluded:
                key_display = key.replace('_', ' ').title()
                if isinstance(value, list):
                    value = ', '.join(str(v) for v in value)
                lines.append(f"**{key_display}:** {value}  ")

        lines.append("\n---\n")
        return '\n'.join(lines)

    def convert_to_docx(
        self,
        content: str,
        output_path: str,
        include_front_matter: bool = False
    ) -> str:
        """
        Convert markdown to Word document.

        Args:
            content: Markdown content with YAML front matter
            output_path: Path for output .docx file
            include_front_matter: Include front matter in document header

        Returns:
            Path to generated document

        Raises:
            ConversionError: If conversion fails
            IOError: If file cannot be written

        Example:
            >>> path = converter.convert_to_docx(
            ...     markdown_string,
            ...     '/tmp/output.docx'
            ... )
        """
        logger.info(f'Converting to DOCX: {output_path}')

        try:
            # Parse content
            metadata, md_content = self.parse_markdown(content)

            # Convert markdown body (without front matter in body)
            # Prepare Pandoc arguments
            extra_args = [
                '--standalone',
                '--highlight-style=pygments',
            ]

            # Add template if provided
            if self.template_path and Path(self.template_path).exists():
                extra_args.append(f'--reference-doc={self.template_path}')
                logger.debug(f'Using template: {self.template_path}')

            # Convert with pypandoc (just the body content)
            pypandoc.convert_text(
                md_content,
                'docx',
                format='md',
                outputfile=output_path,
                extra_args=extra_args
            )

            # Post-process: Fix table formatting (auto-fit columns, add borders)
            self._fix_table_formatting(output_path)

            # Post-process: Add front matter to first page header
            if include_front_matter and metadata:
                self._add_front_matter_to_header(output_path, metadata)
                logger.debug('Added front matter to document header')

            file_size = Path(output_path).stat().st_size
            logger.info(f'Successfully created DOCX: {output_path} ({file_size} bytes)')
            return output_path

        except Exception as e:
            error_msg = f'DOCX conversion failed: {str(e)}'
            logger.error(error_msg, exc_info=True)
            raise ConversionError(error_msg) from e

    def convert_to_pdf(
        self,
        content: str,
        output_path: str,
        include_front_matter: bool = True,
        css_style: Optional[str] = None
    ) -> str:
        """
        Convert markdown to PDF with page numbers.

        Args:
            content: Markdown content with YAML front matter
            output_path: Path for output .pdf file
            include_front_matter: Include front matter in document
            css_style: Custom CSS for styling (uses default if None)

        Returns:
            Path to generated document

        Raises:
            ConversionError: If conversion fails
            IOError: If file cannot be written
        """
        if not WEASYPRINT_AVAILABLE:
            raise ConversionError(
                "WeasyPrint is not available. PDF conversion requires WeasyPrint "
                "with system dependencies installed."
            )

        logger.info(f'Converting to PDF: {output_path}')

        try:
            # Parse content
            metadata, md_content = self.parse_markdown(content)

            # Convert markdown to HTML
            md_processor = markdown.Markdown(extensions=[
                'extra',        # Tables, fenced code, etc.
                'codehilite',   # Syntax highlighting
                'toc',          # Table of contents
                'nl2br',        # Newline to break
                'sane_lists',   # Better list handling
            ])

            html_body = md_processor.convert(md_content)

            # Build front matter HTML
            front_matter_html = ""
            if include_front_matter and metadata:
                front_matter_html = self._format_front_matter_html(metadata)

            # Default CSS if none provided
            if css_style is None:
                css_style = self._get_default_css(metadata.get('title', ''))

            # Build complete HTML document
            html_document = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>{metadata.get('title', 'Document')}</title>
                <style>{css_style}</style>
            </head>
            <body>
                {front_matter_html}
                {html_body}
            </body>
            </html>
            """

            # Generate PDF with WeasyPrint
            HTML(string=html_document).write_pdf(output_path)

            file_size = Path(output_path).stat().st_size
            logger.info(f'Successfully created PDF: {output_path} ({file_size} bytes)')
            return output_path

        except Exception as e:
            error_msg = f'PDF conversion failed: {str(e)}'
            logger.error(error_msg, exc_info=True)
            raise ConversionError(error_msg) from e

    def convert_to_both(
        self,
        content: str,
        base_name: str,
        output_dir: str = '.'
    ) -> Tuple[str, str]:
        """
        Convert markdown to both DOCX and PDF.

        Args:
            content: Markdown content with YAML front matter
            base_name: Base name for output files (no extension)
            output_dir: Directory for output files

        Returns:
            Tuple of (docx_path, pdf_path)

        Example:
            >>> docx, pdf = converter.convert_to_both(
            ...     markdown_string,
            ...     'document',
            ...     '/tmp/output'
            ... )
        """
        logger.info(f'Converting to both formats: {base_name}')

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        docx_path = str(output_dir / f"{base_name}.docx")
        pdf_path = str(output_dir / f"{base_name}.pdf")

        self.convert_to_docx(content, docx_path, include_front_matter=True)
        self.convert_to_pdf(content, pdf_path)

        logger.info(f'Successfully converted to both formats')
        return docx_path, pdf_path

    def _fix_table_formatting(self, docx_path: str) -> None:
        """
        Fix table formatting in DOCX to use auto-fit and proper borders.

        Args:
            docx_path: Path to the DOCX file to modify

        This function:
        - Sets tables to auto-fit content (evenly distributed columns)
        - Adds borders around all table cells
        """
        try:
            doc = Document(docx_path)

            if not doc.tables:
                logger.debug('No tables found in document')
                return

            for table in doc.tables:
                tbl = table._element
                tblPr = tbl.tblPr

                # Remove existing width specifications on columns
                # This allows Word to auto-calculate based on content
                for row in table.rows:
                    for cell in row.cells:
                        # Remove cell width
                        tcPr = cell._element.tcPr
                        if tcPr is not None:
                            # Remove tcW (cell width) elements
                            for tcW in tcPr.findall(qn('w:tcW')):
                                tcPr.remove(tcW)

                # Set table layout to auto (auto-fit content)
                tblLayout = tblPr.find(qn('w:tblLayout'))
                if tblLayout is None:
                    tblLayout = OxmlElement('w:tblLayout')
                    tblPr.append(tblLayout)
                tblLayout.set(qn('w:type'), 'autofit')

                # Set table width to 100% of page width
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is None:
                    tblW = OxmlElement('w:tblW')
                    tblPr.append(tblW)
                tblW.set(qn('w:w'), '5000')
                tblW.set(qn('w:type'), 'pct')

                # Add table borders
                tblBorders = tblPr.find(qn('w:tblBorders'))
                if tblBorders is None:
                    tblBorders = OxmlElement('w:tblBorders')
                    tblPr.append(tblBorders)

                # Define border style
                border_attrs = {
                    qn('w:val'): 'single',
                    qn('w:sz'): '4',  # Border size (1/8 pt)
                    qn('w:space'): '0',
                    qn('w:color'): '000000'  # Black
                }

                # Add all borders (top, left, bottom, right, insideH, insideV)
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = tblBorders.find(qn(f'w:{border_name}'))
                    if border is None:
                        border = OxmlElement(f'w:{border_name}')
                        tblBorders.append(border)
                    for attr, val in border_attrs.items():
                        border.set(attr, val)

            doc.save(docx_path)
            logger.info(f'Fixed formatting for {len(doc.tables)} table(s)')

        except Exception as e:
            logger.warning(f'Failed to fix table formatting: {e}', exc_info=True)
            # Don't raise - document is still valid, just without fixed tables

    def _add_front_matter_to_header(self, docx_path: str, metadata: Dict) -> None:
        """
        Add front matter to the first page header of a DOCX file.

        Args:
            docx_path: Path to the DOCX file to modify
            metadata: Front matter metadata dict

        Raises:
            Exception: If DOCX modification fails
        """
        try:
            # Open the document
            doc = Document(docx_path)

            # Get or create the first section
            if not doc.sections:
                logger.warning("Document has no sections, cannot add header")
                return

            first_section = doc.sections[0]

            # Enable different first page header
            first_section.different_first_page_header_footer = True

            # Get the first page header
            first_page_header = first_section.first_page_header

            # Clear any existing content in the first page header
            for paragraph in first_page_header.paragraphs:
                for run in paragraph.runs:
                    run.clear()

            # Format and add front matter content
            header_text = self._format_front_matter_for_docx_header(metadata)

            # Add the formatted text to header
            if header_text:
                # Clear default paragraphs if they exist
                while len(first_page_header.paragraphs) > 0:
                    first_page_header.paragraphs[0]._element.getparent().remove(
                        first_page_header.paragraphs[0]._element
                    )

                # Add front matter lines
                for line in header_text.split('\n'):
                    if line:  # Skip empty lines
                        p = first_page_header.add_paragraph(line)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        # Make the text smaller for header
                        for run in p.runs:
                            run.font.size = Pt(10)

            # Save the modified document
            doc.save(docx_path)
            logger.debug(f"Successfully added front matter to header in {docx_path}")

        except Exception as e:
            logger.error(f"Failed to add front matter to header: {e}", exc_info=True)
            # Don't raise - we still have a valid document, just without header
            # This allows the conversion to succeed even if header fails

    def _format_front_matter_for_docx_header(self, metadata: Dict) -> str:
        """
        Format front matter as plain text for DOCX header (similar to Google Docs).

        Args:
            metadata: Front matter key-value pairs

        Returns:
            str: Formatted header text
        """
        lines = []

        # Add title if present
        if 'title' in metadata:
            lines.append(metadata['title'])

        # Add other metadata
        for key, value in metadata.items():
            if key != 'title':
                # Format key nicely
                key_display = key.replace('_', ' ').title()

                # Handle different value types
                if isinstance(value, list):
                    value_str = ', '.join(str(v) for v in value)
                elif hasattr(value, 'strftime'):  # Date object
                    value_str = value.strftime('%Y-%m-%d')
                else:
                    value_str = str(value)

                lines.append(f"{key_display}: {value_str}")

        return '\n'.join(lines)

    def _format_front_matter_html(self, metadata: Dict) -> str:
        """
        Format front matter as HTML for PDF generation.

        Args:
            metadata: Dictionary of front matter fields

        Returns:
            HTML string
        """
        html = '<div class="front-matter">\n'
        html += '<h2>Document Information</h2>\n'

        for key, value in metadata.items():
            key_display = key.replace('_', ' ').title()
            if isinstance(value, list):
                value = ', '.join(str(v) for v in value)
            elif hasattr(value, 'strftime'):
                value = value.strftime('%Y-%m-%d')
            html += f'<p><strong>{key_display}:</strong> {value}</p>\n'

        html += '</div>\n'
        return html

    def _get_default_css(self, title: str = '') -> str:
        """
        Get default CSS styling for PDF generation.

        Args:
            title: Document title for header

        Returns:
            CSS string with page layout and styling
        """
        return f"""
        @page {{
            size: A4;
            margin: 1in 0.75in;

            @bottom-center {{
                content: "Page " counter(page) " of " counter(pages);
                font-size: 10pt;
                color: #666;
                font-family: 'Helvetica', sans-serif;
            }}

            @top-right {{
                content: "{title}";
                font-size: 9pt;
                color: #999;
                font-family: 'Helvetica', sans-serif;
            }}
        }}

        @page :first {{
            @bottom-center {{
                content: "";
            }}
            @top-right {{
                content: "";
            }}
        }}

        body {{
            font-family: 'Georgia', serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #333;
        }}

        h1, h2, h3, h4, h5, h6 {{
            font-family: 'Helvetica', sans-serif;
            color: #000;
            page-break-after: avoid;
        }}

        h1 {{
            font-size: 24pt;
            margin-top: 0;
            border-bottom: 2px solid #333;
            padding-bottom: 10pt;
        }}

        h2 {{
            font-size: 18pt;
            margin-top: 20pt;
            border-bottom: 1px solid #999;
            padding-bottom: 5pt;
        }}

        h3 {{
            font-size: 14pt;
            margin-top: 15pt;
        }}

        p {{
            margin: 0.5em 0;
            text-align: justify;
        }}

        code {{
            background-color: #f4f4f4;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
            font-size: 9pt;
            color: #c7254e;
        }}

        pre {{
            background-color: #f4f4f4;
            padding: 10pt;
            border-left: 3px solid #666;
            overflow-x: auto;
            page-break-inside: avoid;
            margin: 1em 0;
        }}

        pre code {{
            background-color: transparent;
            padding: 0;
            color: inherit;
        }}

        blockquote {{
            border-left: 4px solid #ddd;
            padding-left: 15pt;
            margin-left: 0;
            color: #666;
            font-style: italic;
            page-break-inside: avoid;
        }}

        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 15pt 0;
            page-break-inside: avoid;
        }}

        th, td {{
            border: 1px solid #ddd;
            padding: 8pt;
            text-align: left;
        }}

        th {{
            background-color: #f4f4f4;
            font-weight: bold;
            color: #000;
        }}

        tr:nth-child(even) {{
            background-color: #fafafa;
        }}

        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1em auto;
        }}

        a {{
            color: #0066cc;
            text-decoration: none;
        }}

        a:hover {{
            text-decoration: underline;
        }}

        ul, ol {{
            margin: 0.5em 0;
            padding-left: 2em;
        }}

        li {{
            margin: 0.25em 0;
        }}

        .front-matter {{
            border: 2px solid #333;
            background-color: #f9f9f9;
            padding: 15pt;
            margin-bottom: 20pt;
            page-break-after: avoid;
        }}

        .front-matter h2 {{
            margin-top: 0;
            font-size: 14pt;
            border-bottom: none;
        }}

        .front-matter p {{
            margin: 0.25em 0;
            text-align: left;
        }}

        hr {{
            border: none;
            border-top: 1px solid #ccc;
            margin: 2em 0;
        }}
        """
