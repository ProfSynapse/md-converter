"""
Template Generator
Location: /Users/jrosenbaum/Documents/Code/md-converter/app/utils/template_generator.py

This module creates a Word template with page numbers in the footer.
The template is used by pandoc to generate Word documents with automatic page numbering.

Dependencies:
    - python-docx: Word document manipulation
    - docx.oxml: Low-level XML manipulation for page numbers

Used by: app/__init__.py during application initialization
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging
from pathlib import Path


logger = logging.getLogger(__name__)


def create_page_number_element():
    """
    Create the XML element for page numbers in Word format "Page X of Y".

    Returns:
        OxmlElement: XML element containing page number field codes

    Technical details:
        Uses Word field codes:
        - PAGE: Current page number
        - NUMPAGES: Total page count
    """
    # Create a run element with field codes for page numbers
    run = OxmlElement('w:r')

    # Add "Page " text
    t1 = OxmlElement('w:t')
    t1.set(qn('xml:space'), 'preserve')
    t1.text = 'Page '
    run.append(t1)

    # Add PAGE field (current page)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run.append(fldChar1)

    instrText1 = OxmlElement('w:instrText')
    instrText1.set(qn('xml:space'), 'preserve')
    instrText1.text = "PAGE"
    run.append(instrText1)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run.append(fldChar2)

    # Add " of " text
    t2 = OxmlElement('w:t')
    t2.set(qn('xml:space'), 'preserve')
    t2.text = ' of '
    run.append(t2)

    # Add NUMPAGES field (total pages)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run.append(fldChar3)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = "NUMPAGES"
    run.append(instrText2)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run.append(fldChar4)

    return run


def add_table_autofit_style(doc: Document) -> None:
    """
    Add table auto-fit settings to the document's table style.

    This ensures tables in converted documents automatically resize columns
    to fit content rather than using fixed widths.

    Args:
        doc: Document object to modify
    """
    # Create a sample table with auto-fit properties
    # Pandoc will copy these properties when creating tables
    table = doc.add_table(rows=1, cols=3)

    # Set table to auto-fit content
    tbl = table._element
    tblPr = tbl.tblPr

    # Add table layout property for auto-fit
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'auto')
    tblPr.append(tblLayout)

    # Set table width to 100% of page width
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')  # Percentage
    tblPr.append(tblW)

    # Remove the sample table (we just needed it to set the style)
    doc._element.body.remove(table._element)

    logger.debug('Added table auto-fit style to template')


def create_word_template_with_page_numbers(output_path: str) -> str:
    """
    Create a Word template document with page numbers in the footer and auto-fit tables.

    Args:
        output_path: Path where the template will be saved

    Returns:
        Absolute path to the created template

    Raises:
        IOError: If template cannot be written

    Example:
        >>> template_path = create_word_template_with_page_numbers('/app/templates/default.docx')
        >>> converter = MarkdownConverter(template_path=template_path)
    """
    logger.info(f'Creating Word template with page numbers: {output_path}')

    try:
        # Create new document
        doc = Document()

        # Access the footer of the default section
        section = doc.sections[0]
        footer = section.footer

        # Create a paragraph in the footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number element to paragraph
        paragraph._element.append(create_page_number_element())

        # Style the page number text
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

        # Set page margins (standard 1 inch)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Add table auto-fit settings
        add_table_autofit_style(doc)

        # Ensure output directory exists
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)

        # Save template
        doc.save(output_path)

        file_size = Path(output_path).stat().st_size
        logger.info(f'Successfully created template: {output_path} ({file_size} bytes)')

        return str(Path(output_path).resolve())

    except Exception as e:
        error_msg = f'Failed to create Word template: {str(e)}'
        logger.error(error_msg, exc_info=True)
        raise IOError(error_msg) from e


def ensure_template_exists(template_path: str) -> str:
    """
    Ensure the Word template exists, creating it if necessary.

    Args:
        template_path: Path where template should exist

    Returns:
        Absolute path to the template

    Example:
        >>> template = ensure_template_exists('/app/templates/default.docx')
        >>> converter = MarkdownConverter(template_path=template)
    """
    template_file = Path(template_path)

    if template_file.exists():
        logger.debug(f'Template already exists: {template_path}')
        return str(template_file.resolve())

    logger.info(f'Template not found, creating: {template_path}')
    return create_word_template_with_page_numbers(template_path)
